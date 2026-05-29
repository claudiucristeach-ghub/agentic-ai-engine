"""Handler for agent team management and orchestration."""

from collections.abc import AsyncGenerator
from io import BytesIO

import structlog

from google.adk.runners import Runner
from google.genai.types import Content, Part

from app import config
from app.agent_repo import get_agent
from app.handlers.agent_team_app import create_agent_app
from app.handlers.session_handler import SessionHandler
from app.context.artifacts.artifact_service_handler import artifact_service_handler
from app.context.memory.memory_bank_handler import memory_bank_handler

logger = structlog.get_logger(__name__)

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed; cannot extract .docx text")
        return "[Could not read .docx: python-docx is not installed on the server]"

    try:
        doc = Document(BytesIO(data))
    except Exception as e:
        logger.warning("Failed to parse .docx", error=str(e))
        return f"[Could not read .docx file: {e}]"

    lines: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


class AgentTeamHandler:
    def __init__(self, session_handler: SessionHandler) -> None:
        self._session_handler = session_handler
        self.session_id: str | None = None
        self.runner: Runner | None = None
        self.agent_id: str | None = None

    async def create_agent_team_runner(self, agent_id: str) -> str:
        try:
            agent = get_agent(agent_id)
            agent_app = create_agent_app(agent)

            session = await self._session_handler.get_or_create_session(
                app_name=agent_app.name,
                agent_id=agent_id,
            )

            self.session_id = session.id
            self.agent_id = agent_id

            runner_kwargs = dict(
                app=agent_app,
                session_service=self._session_handler.service,
                artifact_service=artifact_service_handler.service,
            )

            if memory_bank_handler.service is not None:
                runner_kwargs["memory_service"] = memory_bank_handler.service

            self.runner = Runner(**runner_kwargs)

            logger.info(
                "Created agent runner",
                agent_id=agent_id,
                session_id=self.session_id,
            )

        except KeyError:
            logger.error("Unknown agent requested", agent_id=agent_id)
            raise ValueError(f"Unknown agent: {agent_id}")
        except Exception as e:
            logger.error(
                "Failed to create agent runner",
                agent_id=agent_id,
                error=str(e),
            )
            raise

        return self.session_id

    async def switch_agent(self, agent_id: str) -> str:
        logger.info("Switching agent", from_agent=self.agent_id, to_agent=agent_id)
        return await self.create_agent_team_runner(agent_id)

    async def create_new_session(self) -> str:
        if not self.agent_id:
            raise ValueError("No agent is currently active.")

        agent = get_agent(self.agent_id)
        agent_app = create_agent_app(agent)

        session = await self._session_handler.reset_session(
            app_name=agent_app.name,
            agent_id=self.agent_id,
        )

        self.session_id = session.id

        runner_kwargs = dict(
            app=agent_app,
            session_service=self._session_handler.service,
            artifact_service=artifact_service_handler.service,
        )

        if memory_bank_handler.service is not None:
            runner_kwargs["memory_service"] = memory_bank_handler.service

        self.runner = Runner(**runner_kwargs)

        logger.info(
            "Created new session",
            agent_id=self.agent_id,
            session_id=self.session_id,
        )

        return self.session_id

    @staticmethod
    def _build_user_content(
        text: str,
        files: list[dict] | None = None,
    ) -> Content:
        parts: list[Part] = []

        text_mime_prefixes = (
            "text/",
            "application/json",
            "application/xml",
            "application/javascript",
            "application/x-yaml",
        )

        if files:
            for f in files:
                mime: str = f.get("mime", "application/octet-stream")
                data: bytes = f["data"]
                name: str = f.get("name", "file")

                if any(mime.startswith(p) for p in text_mime_prefixes):
                    try:
                        decoded = data.decode("utf-8")
                    except UnicodeDecodeError:
                        decoded = data.decode("latin-1")
                    parts.append(Part(text=f"[File: {name}]\n{decoded}"))

                elif mime == _DOCX_MIME or name.lower().endswith(".docx"):
                    extracted = _extract_docx_text(data)
                    parts.append(Part(text=f"[File: {name}]\n{extracted}"))

                else:
                    parts.append(Part(inline_data={"mime_type": mime, "data": data}))

        if text:
            parts.append(Part(text=text))

        if not parts:
            parts.append(Part(text=""))

        return Content(parts=parts, role="user")

    async def get_agent_response(
        self,
        user_query: str,
        *,
        files: list[dict] | None = None,
    ) -> str:
        if not self.runner or not self.session_id:
            raise RuntimeError("Agent team runner has not been created yet.")

        user_message = self._build_user_content(user_query, files)

        try:
            final_response_text = ""

            for event in self.runner.run(
                user_id=config.USER_ID,
                session_id=self.session_id,
                new_message=user_message,
            ):
                if event.is_final_response() and event.content and event.content.parts:
                    final_response_text = "".join(
                        part.text for part in event.content.parts if part.text
                    )

            logger.info(
                "Received response from agent team",
                response_length=len(final_response_text),
            )

            return final_response_text

        except Exception as e:
            logger.error("Failed to get response from agent team", error=str(e))
            raise

    async def stream_agent_response(
        self,
        user_query: str,
        *,
        files: list[dict] | None = None,
    ) -> AsyncGenerator[dict, None]:
        if not self.runner or not self.session_id:
            raise RuntimeError("Agent team runner has not been created yet.")

        user_message = self._build_user_content(user_query, files)

        try:
            async for event in self.runner.run_async(
                user_id=config.USER_ID,
                session_id=self.session_id,
                new_message=user_message,
            ):
                if event.content and event.content.parts:
                    text = "".join(
                        part.text for part in event.content.parts if part.text
                    )
                    if text:
                        yield {
                            "type": "final" if event.is_final_response() else "partial",
                            "author": event.author or "agent",
                            "content": text,
                        }

            if memory_bank_handler.service is not None:
                session = await self._session_handler.service.get_session(
                    app_name=self.runner.app.name,
                    user_id=config.USER_ID,
                    session_id=self.session_id,
                )
                if session is not None:
                    await memory_bank_handler.service.add_session_to_memory(session)
                    logger.info(
                        "Session added to memory bank",
                        agent_id=self.agent_id,
                        session_id=self.session_id,
                    )
                logger.info(
                    "Session added to memory bank",
                    agent_id=self.agent_id,
                    session_id=self.session_id,
                )

        except Exception as e:
            logger.error("Streaming error", error=str(e))
            yield {"type": "error", "author": "system", "content": str(e)}